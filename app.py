import streamlit as st
from langchain.docstore.document import Document
from embed_store import embed_and_store
from rag_chain import build_rag_chain
from web_scrapper import scrape_website

from PyPDF2 import PdfReader
import docx
import datetime
import base64
import os
os.environ["LANGCHAIN_API_KEY"] = "none"


st.set_page_config(page_title="RAG Chatbot")
st.title("RAG - Chat with Website + Documents")

# Session states
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "rag_chain" not in st.session_state:
    st.session_state.rag_chain = None
if "last_url" not in st.session_state:
    st.session_state.last_url = ""

# ✅ File reader for PDF, DOCX, TXT, etc.
def read_file(uploaded_file):
    filename = uploaded_file.name.lower()
    try:
        if filename.endswith(".pdf"):
            reader = PdfReader(uploaded_file)
            text = "".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif filename.endswith(".docx"):
            doc = docx.Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            text = uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"❌ Error reading {filename}: {e}")
        text = ""
    return text

# Save chat history to a local file (optional)
def save_chat_to_file():
    now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"chat_{now}.txt"
    with open(filename, "w", encoding="utf-8") as f:
        for q, a in st.session_state.chat_history:
            f.write(f"You: {q}\nGemma: {a}\n\n")
    return filename

# Sidebar
with st.sidebar:
    st.subheader("📥 Ingest Website + Upload Docs")
    url = st.text_input("🌐 Website URL", placeholder="https://example.com")
    uploaded_files = st.file_uploader("📄 Upload Documents", type=["pdf", "txt", "md", "docx"], accept_multiple_files=True)

    if st.button("🔍 Ingest Content"):
        with st.spinner("Processing documents + scraping website..."):
            try:
                all_text = []
                if url:
                    site_text = scrape_website(url)
                    st.write(f"✅ Scraped {len(site_text)} characters from site")
                    st.session_state.last_url = url
                    all_text.append(Document(page_content=site_text))

                for file in uploaded_files:
                    text = read_file(file)
                    if text:
                        all_text.append(Document(page_content=text))

                if all_text:
                    vectordb, session_dir = embed_and_store(all_text)
                    st.session_state.rag_chain = build_rag_chain(vectordb)
                    st.success("✅ Ingestion complete. You can now ask questions!")
                else:
                    st.error("❌ No readable content found.")
            except Exception as e:
                st.error(f"❌ Error: {e}")

st.divider()

# Chat interface
query = st.text_input("💬 Ask your question here")
if query and st.session_state.rag_chain:
    with st.spinner("💡 Thinking..."):
        try:
            result = st.session_state.rag_chain(query)
            answer = result["result"]
            context_docs = result.get("source_documents", [])
            st.session_state.chat_history.append((query, answer, context_docs))
        except Exception as e:
            st.error(f"❌ Error: {e}")

# Display chat
if st.session_state.chat_history:
    st.subheader("📜 Chat History")
    for q, a, docs in reversed(st.session_state.chat_history):
        st.markdown(f"**🧚‍♀️You:** {q}")
        st.markdown(f"**🤖 Gemma:** {a}")

        # 🔍 Show retrieved context
        #if docs:
        #    with st.expander("📄 Retrieved Context", expanded=False):
        #       for i, doc in enumerate(docs):
        #           content = doc.page_content[:500]
        #           st.markdown(f"**Doc {i+1}:**\n```{content}```")

    # Save chat button
    if st.button("💾 Save Chat to File"):
        filename = save_chat_to_file()
        with open(filename, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
            href = f'<a href="data:file/txt;base64,{b64}" download="{filename}">Click to download</a>'
            st.markdown(href, unsafe_allow_html=True)
